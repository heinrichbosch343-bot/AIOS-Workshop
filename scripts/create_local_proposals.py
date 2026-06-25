"""
create_local_proposals.py  —  generates 7 realistic mock client proposals as
formatted Word (.docx) files in a local folder, ready to drag into Google Drive
for the data-pooling LinkedIn demo.

Run:
    python scripts/create_local_proposals.py

Output:  outputs/demo-proposals/*.docx
No Google / network access required — pure local file generation.
"""

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = Path(__file__).resolve().parents[1] / "outputs" / "demo-proposals"

# Each proposal: a clean filename + the document title + structured body.
PROPOSALS = [
    {
        "filename": "Nexgen Manufacturing - Process Automation.docx",
        "title": "Proposal — Nexgen Manufacturing (AI Process Automation)",
        "body": """CLIENT: Nexgen Manufacturing (Pty) Ltd
CONTACT: Sipho Dlamini, Head of Operations
DATE: 12 May 2026
STATUS: Pending — awaiting sign-off from CFO
VALID UNTIL: 30 June 2026

SCOPE OF WORK
Phase 1 — Process Audit & Automation Mapping (4 weeks). Full audit of current production floor workflows, quality control checks, and shift reporting processes. Identify the top 8 automation candidates.
Phase 2 — AIOS Build & Integration (8 weeks). Build custom AI layer integrating with their existing ERP (SAP B1). Automate shift reports, defect classification, and supplier follow-ups.
Phase 3 — Training & Handover (2 weeks). Onsite training with operations team. Full documentation delivered.

INVESTMENT
Phase 1: R 45,000
Phase 2: R 185,000
Phase 3: R 25,000
Monthly Retainer (post-handover): R 12,500/month
TOTAL PROJECT VALUE: R 255,000 + retainer

PAYMENT TERMS
50% upfront, 25% at Phase 2 kickoff, 25% on completion.

NOTES
Client expressed concern about integration with legacy SAP version (6.0). Heinrich to confirm compatibility before sign-off. Decision expected by 20 June.""",
    },
    {
        "filename": "Hartley Retail Group - Customer Intelligence.docx",
        "title": "Proposal — Hartley Retail Group (Customer Intelligence System)",
        "body": """CLIENT: Hartley Retail Group
CONTACT: Anri Hartley, CEO
DATE: 28 April 2026
STATUS: APPROVED — project starting 1 July 2026
VALID UNTIL: N/A (signed)

SCOPE OF WORK
Build a customer intelligence layer across 12 retail branches. Consolidate POS data, loyalty programme, and in-store foot traffic into a single AI dashboard that generates weekly insights and flags stock anomalies.
Deliverables: Central data pool ingesting from 12 branch POS systems. Weekly AI-generated brief delivered to Anri every Monday 7am. Stock anomaly alerts via WhatsApp. Competitor pricing monitor (web scrape + weekly summary).

INVESTMENT
Setup & Build: R 320,000
Monthly Retainer: R 18,000/month (12-month minimum)

PAYMENT TERMS
R 80,000 upfront, balance billed in 4 monthly instalments. Contract start 1 July 2026. Contract length 12 months minimum.

NOTES
Contract signed 15 May 2026. Kick-off meeting scheduled 25 June 2026. Anri wants the first weekly brief live before the school holiday trading period.""",
    },
    {
        "filename": "BrightPath Logistics - Fleet Intelligence.docx",
        "title": "Proposal — BrightPath Logistics (Fleet & Driver Intelligence)",
        "body": """CLIENT: BrightPath Logistics CC
CONTACT: Trevor Mokoena, Operations Director
DATE: 3 June 2026
STATUS: In negotiation — price sensitivity flagged
VALID UNTIL: 15 July 2026

SCOPE OF WORK
AIOS layer to monitor 47-vehicle fleet: automated driver scorecards, route optimisation summaries, and fuel anomaly detection. Weekly AI brief to Trevor and daily SMS to each driver with their score.
Phase 1 — Telematics data integration (Ctrack API): 3 weeks.
Phase 2 — Scoring model + dashboard build: 5 weeks.
Phase 3 — Driver comms automation (WhatsApp/SMS): 2 weeks.

INVESTMENT
Original quote — Build: R 195,000, Monthly: R 9,500/month.
Revised quote (after negotiation, 10 June) — Build: R 165,000, Monthly: R 8,000/month. Condition: 18-month retainer commitment.

NOTES
Trevor pushed back on original price — he has a competing quote from a local dev shop at R 140k build-only (no AI, no retainer model). Heinrich to send a value comparison doc by 18 June. Decision expected EOD 25 June.""",
    },
    {
        "filename": "Sandstone Medical Group - HR Compliance.docx",
        "title": "Proposal — Sandstone Medical Group (HR & Compliance Automation)",
        "body": """CLIENT: Sandstone Medical Group
CONTACT: Dr. Fatima Essop, Practice Manager
DATE: 18 March 2026
STATUS: CLOSED — client went with internal IT team
VALID UNTIL: Expired

SCOPE OF WORK
Automate HR onboarding for clinical and admin staff, compliance document tracking (HPCSA renewals, BLS certifications), and monthly payroll summary generation across 3 practices.

INVESTMENT
Build: R 145,000
Monthly: R 7,500/month

OUTCOME
Client decided to build internally using a junior developer. Decision made 4 May 2026. Relationship maintained — Dr. Essop open to revisiting in 12 months if internal build stalls (she flagged this explicitly).

LESSON LEARNED
Medical clients have longer decision cycles due to committee buy-in requirements. Follow up with Sandstone in November 2026.""",
    },
    {
        "filename": "Cape Agri Cooperative - Demand Forecasting.docx",
        "title": "Proposal — Cape Agri Cooperative (Seasonal Demand Forecasting)",
        "body": """CLIENT: Cape Agri Cooperative
CONTACT: Pieter van Zyl, General Manager
DATE: 20 May 2026
STATUS: Pending — follow-up required (gone quiet since 1 June)
VALID UNTIL: 20 July 2026

SCOPE OF WORK
Build a seasonal demand forecasting tool fed by 3 years of historical sales data, weather API, and regional crop cycle data. Output: a weekly planting and stock recommendation for the co-op's 220 member farmers.
Deliverables: Data ingestion pipeline (Excel exports + weather API). Forecasting model (AI-driven, retrained monthly). Member-facing weekly WhatsApp summary. GM dashboard with override controls.

INVESTMENT
Build: R 210,000
Monthly: R 11,000/month

PAYMENT TERMS
Milestone-based — 4 equal payments of R 52,500.

NOTES
Pieter was enthusiastic in the first meeting but has not responded to follow-up emails sent 3 June and 10 June. Possible budget freeze due to drought conditions affecting member revenues. Heinrich to call directly before 25 June. Risk: HIGH.""",
    },
    {
        "filename": "Elevate Property Group - Investment Intelligence.docx",
        "title": "Proposal — Elevate Property Group (Investment Intelligence Dashboard)",
        "body": """CLIENT: Elevate Property Group
CONTACT: Nomsa Khumalo, Investment Director
DATE: 7 June 2026
STATUS: Proposal sent — first response expected week of 23 June
VALID UNTIL: 31 July 2026

SCOPE OF WORK
Build an investment intelligence layer across Elevate's portfolio of 34 commercial properties. Consolidate rental income, vacancy rates, maintenance costs, and market comparables into a single AI-powered dashboard. Monthly AI brief to the board.
Key features: Portfolio health score (AI-generated, updated monthly). Automated tenant risk flags (missed payments, upcoming renewals). Market comparable tracker (auto-pulls from PropData API). Board-ready PDF brief generated on the 1st of each month.

INVESTMENT
Build: R 285,000
Monthly Retainer: R 15,500/month

PAYMENT TERMS
40% upfront (R 114,000), 60% on go-live. Estimated go-live: 10 weeks from contract signature.

NOTES
Nomsa was referred by Anri Hartley (Hartley Retail Group — existing client). Warm lead. She reviewed our work for Hartley before the meeting. Strong fit.""",
    },
    {
        "filename": "TechBridge Academy - Student Performance AI.docx",
        "title": "Proposal — TechBridge Academy (Student Performance AI)",
        "body": """CLIENT: TechBridge Academy (Private School Group — 4 campuses)
CONTACT: Mr. Deon Fredericks, Executive Principal
DATE: 1 June 2026
STATUS: Pending — board presentation scheduled 30 June 2026
VALID UNTIL: 31 August 2026

SCOPE OF WORK
Phase 1 — Data Consolidation (6 weeks). Ingest student performance data from 4 campuses (currently in 3 different SIS platforms). Normalise and pool into a central data store.
Phase 2 — Intelligence Layer (6 weeks). AI model to flag at-risk students (attendance + performance trends). Weekly brief to each campus head. Monthly board summary.
Phase 3 — Parent Communication Automation (4 weeks). Auto-generate personalised progress reports and flag parents when intervention is recommended.

INVESTMENT
Build (all 3 phases): R 380,000
Monthly Retainer: R 16,000/month

PAYMENT TERMS
3 milestone payments — R 127,000 per phase completion.

NOTES
Board is conservative — Deon believes they'll approve if Heinrich can present ROI in terms of student retention (each student = ~R 85,000/year in fees). Heinrich to prepare a 1-page ROI case before the 30 June board meeting. Decision by 15 July 2026.""",
    },
]

# Lines that act as section headers are short and ALL CAPS.
_HEADER_RE = re.compile(r"^[A-Z][A-Z &/]{2,}$")
# "KEY: value" lines at the top of each doc.
_KEYVAL_RE = re.compile(r"^([A-Z][A-Za-z ]+):\s*(.+)$")

ACCENT = RGBColor(0x1D, 0x4E, 0xD8)   # AIOS blue
MUTED = RGBColor(0x55, 0x55, 0x55)


def build_doc(title: str, body: str) -> Document:
    doc = Document()

    # Brand strip
    brand = doc.add_paragraph()
    run = brand.add_run("BOSCHAI  ·  AIOS DELIVERY")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = ACCENT

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for raw in body.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if _HEADER_RE.match(line):
            doc.add_heading(line.title(), level=2)
            continue
        m = _KEYVAL_RE.match(line)
        if m and len(m.group(1)) < 16:
            p = doc.add_paragraph()
            k = p.add_run(f"{m.group(1)}:  ")
            k.bold = True
            k.font.color.rgb = MUTED
            p.add_run(m.group(2))
        else:
            doc.add_paragraph(line)

    # Footer
    foot = doc.add_paragraph()
    fr = foot.add_run("Prepared by Heinrich · BoschAI · Confidential")
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = MUTED
    return doc


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Writing {len(PROPOSALS)} proposals to:\n  {OUT_DIR}\n")
    for i, p in enumerate(PROPOSALS, 1):
        doc = build_doc(p["title"], p["body"])
        out_path = OUT_DIR / p["filename"]
        doc.save(str(out_path))
        print(f"  [{i}/{len(PROPOSALS)}] {p['filename']}")
    print(f"\nDone. {len(PROPOSALS)} files created.")
    print("Next: drag them into a Google Drive folder, then pick that folder in the dashboard.")


if __name__ == "__main__":
    main()
